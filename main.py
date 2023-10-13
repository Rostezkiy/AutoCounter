import locale
from datetime import datetime

from docx import Document

try:
    doc = Document('Отчет.docx')
    locale.setlocale(locale.LC_ALL, 'ru_RU.UTF-8')
    date = datetime.now().strftime('«%d» %B %Yг')
    paragraphs = doc.paragraphs
    table = doc.tables[0]

    for i in range(len(table.columns)):
        table.cell(1, i).text = table.cell(2, i).text

    table.cell(2, 0).text = datetime.now().strftime('%d %B')
    table.cell(2, 1).text = input('"Холодная вода куб. м.": ') + " куб. м."
    table.cell(2, 2).text = input('"Горячая вода куб. м.": ') + " куб. м."

    cold_water_1 = float(table.cell(1, 1).text.split()[0])
    hot_water_1 = float(table.cell(1, 2).text.split()[0])
    cold_water_2 = float(table.cell(2, 1).text.split()[0])
    hot_water_2 = float(table.cell(2, 2).text.split()[0])

    table.cell(3, 1).text = f'{round(cold_water_2 - cold_water_1, 2)} куб. м.'
    table.cell(3, 2).text = f'{round(hot_water_2 - hot_water_1, 2)} куб. м.'

    paragraphs[-2].text = f'Подпись квартиросъемщика_____________________________________Дата {date}'
    paragraphs[-1].text = f'Сведения принял______________________________________________Дата {date}'

    doc.save(f"Отчет_{datetime.now().strftime('%d %B %Yг')}.docx")


except:
    print("Поместите файл с названием \"Отчет.docx\" рядом с исполняемым файлом")
    input()
